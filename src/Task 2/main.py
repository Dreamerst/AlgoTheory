import coat
import three_piece_suit
import trousers
import docx
import openpyxl

clothes = []
number = 0
while True:
    print("Выберите действие:")
    print("1. Добавить пиджак")
    print("2. Добавить брюки")
    print("3. Добавить костюм-тройку")
    print("4. Посчитать количество ткани на выбранную одежду")
    print("5. Посчитать стоимость выбранной одежды")
    print("6. Сохранить в .doc")
    print("7. Сохранить в .xls")
    print("8. Выход")

    number = int(input())
    match number:
        case 1:
            print("Введите размер пиджака:")
            size = int(input())
            print("Введите количество пуговиц:")
            buttons = int(input())
            my_coat = coat.Coat(size, buttons)
            clothes.append(my_coat)
            print(f"Необходимое количество ткани: {my_coat.fabric_calculation()}")
            print(f"Стоимость пиджака: {my_coat.cost_calculation()} руб.")
        case 2:
            print("Введите размер брюк:")
            size = int(input())
            print("Добавить ремень: y/n")
            belt = input()
            if belt == "y":
                belt_bool = True
            elif belt == "n":
                belt_bool = False
            else:
                raise ValueError("You should choose only 'y' or 'n'")
            my_trousers = trousers.Trousers(size, belt_bool)
            clothes.append(my_trousers)
            print(f"Необходимое количество ткани: {my_trousers.fabric_calculation()}")
            print(f"Стоимость брюк: {my_trousers.cost_calculation()} руб.")
        case 3:
            print("Введите размер костюма:")
            size = int(input())
            print("Введите количество пуговиц пиджака:")
            coat_buttons = int(input())
            print("Введите количество пуговиц жилета:")
            jacket_buttons = int(input())
            print("Добавить ремень: y/n")
            belt = input()
            if belt == "y":
                belt_bool = True
            elif belt == "n":
                belt_bool = False
            else:
                raise ValueError("You should choose only 'y' or 'n'")
            my_three_piece_suit = three_piece_suit.ThreePieceSuit(size, coat_buttons, jacket_buttons, belt_bool)
            clothes.append(my_three_piece_suit)
            print(f"Необходимое количество ткани: {my_three_piece_suit.fabric_calculation()}")
            print(f"Стоимость костюма-троки: {my_three_piece_suit.cost_calculation()} руб.")
        case 4:
            fabric = 0
            for suit in clothes:
                fabric += suit.fabric_calculation()
            print(f"Общий расход ткани: {fabric}")
        case 5:
            cost = 0
            for suit in clothes:
                cost += suit.cost_calculation()
            print(f"Общая стоимость: {cost}")
        case 6:
            doc = docx.Document()
            fabric = 0
            cost = 0
            for suit in clothes:
                fabric += suit.fabric_calculation()
                cost += suit.cost_calculation()
                doc.add_paragraph(f"{suit.info()}")
                doc.add_paragraph(f"Ткань: {suit.fabric_calculation()}")
                doc.add_paragraph(f"Стоимость: {suit.cost_calculation()} руб.")
            doc.add_paragraph(f"Общий расход ткани: {fabric}")
            doc.add_paragraph(f"Общая стоимость: {cost}")
            doc.save("result.doc")
        case 7:
            doc = openpyxl.Workbook()
            sheet = doc.active

            cell1 = sheet.cell(row=1, column=1)
            cell1.value = "Наименование"
            cell2 = sheet.cell(row=1, column=2)
            cell2.value = "Расход ткани"
            cell3 = sheet.cell(row=1, column=3)
            cell3.value = "Стоимость"

            fabric = 0
            cost = 0
            for i in range(len(clothes)):
                fabric += clothes[i].fabric_calculation()
                cost += clothes[i].cost_calculation()
                cell1 = sheet.cell(row=i + 2, column=1)
                cell1.value = clothes[i].info()
                cell2 = sheet.cell(row=i + 2, column=2)
                cell2.value = clothes[i].fabric_calculation()
                cell3 = sheet.cell(row=i + 2, column=3)
                cell3.value = clothes[i].cost_calculation()
            fabric_cell = sheet.cell(row=len(clothes) + 2, column=2)
            fabric_cell.value = fabric
            cost_cell = sheet.cell(row=len(clothes) + 2, column=3)
            cost_cell.value = cost
            doc.save("result.xls")
        case 8:
            break
        case _:
            continue
